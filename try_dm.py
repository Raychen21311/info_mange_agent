
import streamlit as st
import os
import time
import fitz
import google.generativeai as genai
import pickle
from langchain_community.vectorstores import FAISS
from langchain.docstore.document import Document
from dotenv import load_dotenv
from langchain_huggingface import HuggingFaceEmbeddings
    
from docx import Document as DocxDocument
import pandas as pd
#from pptx import Presentation
from PyPDF2 import PdfReader
import torch
import faiss
import requests
from bs4 import BeautifulSoup
# --- Gemini 設定 ---
load_dotenv()
if os.getenv('GOOGLE_API_KEY'):
    genai.configure(api_key=os.getenv('GOOGLE_API_KEY'))
gen_model = genai.GenerativeModel("gemini-2.5-flash")
#device = "cuda" if torch.cuda.is_available() else "cpu"

# --- 向量模型 ---
embedding_model = HuggingFaceEmbeddings(
    model_name="sentence-transformers/all-mpnet-base-v2"#,
  #  model_kwargs={"device": device}  # ✅ 指定設備
)
# --- 文件讀取與切段落（核心知識庫） ---

# 後台設定網址清單
web_urls = [
    "https://dep.mohw.gov.tw/DOIM/mp-114.html",  # 這裡換成你的實際網址
    "https://www.facebook.com/chnchng.lee.5/?locale=zh_TW"
]



def fetch_webpage_text(url):
    try:
        r = requests.get(url, timeout=10)
        r.raise_for_status()
        soup = BeautifulSoup(r.text, "html.parser")
        return soup.get_text(separator="\n")
    except Exception as e:
        return f"❌ 無法擷取 {url}：{e}"

from langchain.docstore.document import Document

def build_web_vector_store(urls):
    docs = []
    for url in urls:
        text = fetch_webpage_text(url)
        chunks = [text[i:i+1000] for i in range(0, len(text), 1000)]  # 切段落
        docs.extend([Document(page_content=c) for c in chunks])
    return FAISS.from_documents(docs, embedding_model)

# 建立網頁臨時向量庫
web_vector_store = build_web_vector_store(web_urls)

# --- 建立向量資料庫 (FAISS) — AES‑GCM 解密 + 列出檔案清單 ---
import io, zipfile, tempfile, base64, time
from pathlib import Path
from cryptography.hazmat.primitives.ciphers.aead import AESGCM
from langchain_community.vectorstores import FAISS

start = time.perf_counter()

# 直接指定加密檔路徑（相對於 try_dm.py）
enc_path = Path("faiss_index.enc")

def _decrypt_bytes(blob: bytes, key_b64: str) -> bytes:
    key = base64.urlsafe_b64decode(key_b64)  # 32 bytes -> AES‑256‑GCM
    nonce, ct = blob[:12], blob[12:]
    return AESGCM(key).decrypt(nonce, ct, associated_data=None)

try:
    if enc_path.exists():
        with open(enc_path, "rb") as f:
            blob = f.read()
        key_b64 = st.secrets["FAISS_KEY_B64"]  # 金鑰請放在 Secrets
        zip_bytes = _decrypt_bytes(blob, key_b64)

        # 解壓 ZIP 到暫存資料夾
        tmp_dir = Path(tempfile.mkdtemp(prefix="faiss_"))
        with zipfile.ZipFile(io.BytesIO(zip_bytes), "r") as zf:
            zf.extractall(tmp_dir)

        # ✅ 列出解壓後檔案清單（用 print）
        print("📂 解壓後檔案清單：")
        for p in tmp_dir.rglob("*"):
            print(" -", p.relative_to(tmp_dir))

        # 嘗試直接載入（假設 ZIP 根目錄就有 index.faiss 和 index.pkl）
        vector_store = FAISS.load_local(
            str(tmp_dir),
            embeddings=embedding_model,
            allow_dangerous_deserialization=True
        )
        st.write(f"✅ FAISS 載入完成，耗時: {time.perf_counter() - start:.2f} 秒")
    else:
        raise FileNotFoundError(f"找不到加密檔：{enc_path}")

except Exception as e:
    st.error(f"❌ 載入 FAISS 失敗：{e}")
    raise

           






# --- RAG 檢索 ---
def retrieve_context(question, top_k=3):
    start = time.perf_counter()
    docs = vector_store.similarity_search(question, k=top_k)
    docs_web = web_vector_store.similarity_search(question, k=top_k)
    search_time = time.perf_counter() - start
    st.write(f"✅ 檢索耗時: {search_time:.2f} 秒")
    return [d.page_content for d in docs + docs_web]

# --- Gemini 回答生成（加入歷史對話） ---
prompt_info = st.secrets["prompt_info"]
def get_gemini_response(user_question, context_text, history):
    history_text = ""
    for role, msg in history:
        history_text += f"{role}: {msg}\n"
        
    prompt = prompt_info.format(
        user_question=user_question,
        context_text=context_text,
        history_text=history_text
    )


    start = time.perf_counter()
    response = gen_model.generate_content(prompt)
    gemini_time = time.perf_counter() - start
    st.write(f"✅ Gemini API 耗時: {gemini_time:.2f} 秒")
    return response.text


import streamlit as st
import fitz

# --- 初始化 ---
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "user_file_text" not in st.session_state:
    st.session_state.user_file_text = ""

st.set_page_config(page_title="資訊處 AI Agent", layout="wide")

# --- 自訂 CSS ---
st.markdown("""
<style>
.bottom-container {
    position: fixed;
    bottom: 0;
    left: 0;
    right: 0;
    padding: 12px;
    background-color: white;
    border-top: 1px solid #ddd;
    display: flex;
    align-items: center;
    gap: 8px;
}
.input-wrapper {
    flex: 1;
    position: relative;
}
.input-box {
    width: 100%;
    padding: 12px 40px 12px 12px;
    border: 1px solid #ccc;
    border-radius: 20px;
    font-size: 15px;
}
.send-btn {
    position: absolute;
    right: 6px;
    top: 10%;
    transform: translateY(-50%);
    background-color: black;
    color: white;
    border: none;
    border-radius: 50%;
    width: 28px;
    height: 28px;
    font-size: 14px;
    cursor: pointer;
}
.send-btn:hover {
    background-color: #333;
}
</style>
""", unsafe_allow_html=True)

st.title("資訊處 AI Agent")

# --- 聊天顯示區 ---
chat_container = st.container()
for role, msg in st.session_state.chat_history:
    with chat_container:
        with st.chat_message("user" if role == "使用者" else "assistant"):
            st.markdown(msg)

# --- 底部輸入框 ---
st.markdown('<div class="bottom-container">', unsafe_allow_html=True)

with st.form("chat_form", clear_on_submit=True):
    cols = st.columns([2, 8, 0.6])

    # 上傳檔案
    with cols[0]:
        uploaded_files = st.file_uploader(
            "📄 文件",
            type=["pdf", "docx","doc","txt","xls","xlsx","ppt","pptx"],
            key="file_uploader",
            label_visibility="collapsed",
            accept_multiple_files=True
        )
        if uploaded_files:
            for uploaded_file in uploaded_files:  # ✅ 迭代每個檔案
                ext = uploaded_file.name.lower().split(".")[-1]
                file_text = ""

                if ext == "pdf":
                    pdf_document = fitz.open(stream=uploaded_file.read(), filetype="pdf")
                    for page in pdf_document:
                        file_text += page.get_text() + "\n\n"

                elif ext in ["docx", "doc"]:
                    doc = DocxDocument(uploaded_file)
                    file_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

                elif ext == "txt":
                    file_text = uploaded_file.read().decode("utf-8")

                elif ext in ["xls", "xlsx"]:
                    df = pd.read_excel(uploaded_file)
                    file_text = df.to_csv(index=False)

                elif ext in ["ppt", "pptx"]:
                    prs = Presentation(uploaded_file)
                    for slide in prs.slides:
                        for shape in slide.shapes:
                            if hasattr(shape, "text"):
                                file_text += shape.text + "\n"
            
            
            
                st.session_state.user_file_text = file_text
                st.success(f"{uploaded_file.name} 已成功讀取")
  


    # 輸入框 + 送出按鈕
    with cols[1]:
        user_input = st.text_area("請輸入您的問題…", key="chat_input", label_visibility="collapsed",height=50)
        
    with cols[2]:
        submit = st.form_submit_button("➤", use_container_width=False)


st.markdown('</div>', unsafe_allow_html=True)

# --- 處理訊息 ---
if submit and user_input:
    # 1️⃣ 顯示使用者訊息
    with chat_container:
        with st.chat_message("user"):
            st.markdown(user_input)

    # 2️⃣ 在聊天區建立 AI 回答的占位容器
    with chat_container:
        answer_placeholder = st.empty()  # 占位
        with answer_placeholder.container():
            # 在占位容器裡先顯示 spinner
            with st.chat_message("assistant"):
                st.markdown(
            """
            <div style="
                background-color:transparent;  /* 你想要的顏色 */
                padding:8px 12px;
                border-radius:12px;
                display:inline-block;
            ">
                ⏳ 分析中...
            </div>
            """,
            unsafe_allow_html=True
        )
    
    # AI 分析
    context_paragraphs = retrieve_context(user_input)
    context_text = "\n\n".join(context_paragraphs)
    if st.session_state.user_file_text:
        context_text += "\n\n---使用者上傳的附件內容---\n\n" + st.session_state.user_file_text

    answer = get_gemini_response(user_input, context_text, st.session_state.chat_history)

    # 4️⃣ 將占位 spinner 替換成最終回答
    answer_placeholder.empty()  # 清除 spinner
    with chat_container:
        with st.chat_message("assistant"):
            st.markdown(answer)
    # 更新歷史對話
    st.session_state.chat_history.append(("使用者", user_input))
    st.session_state.chat_history.append(("Agent", answer))


  
